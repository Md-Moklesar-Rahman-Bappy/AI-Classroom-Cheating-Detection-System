#!/usr/bin/env python3
"""
Generate AI_Classroom_Cheating_Detection_System_Complete_Documentation.docx
From docs/COMPLETE_PROJECT_DOCUMENTATION.md (intermediate source).
Requires: python-docx
Usage: python scripts/generate_complete_documentation.py
Output: AI_Classroom_Cheating_Detection_System_Complete_Documentation.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_custom(doc, text, level):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(23, 54, 93) if level <= 2 else RGBColor(15, 107, 120)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(36, 52, 71)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    # light background via shading on paragraph
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F7FA')
    p._element.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(36, 52, 71)
    return p


def add_warning_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    # light amber background
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FFF8E1')
    p._element.get_or_add_pPr().append(shading)
    run = p.add_run('⚠  ' + text)
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(183, 121, 31)
    run.font.bold = True
    return p


def add_important_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'E8F4FD')
    p._element.get_or_add_pPr().append(shading)
    run = p.add_run('ℹ  ' + text)
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(37, 99, 235)
    run.font.bold = True
    return p


def add_security_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FCEAEA')
    p._element.get_or_add_pPr().append(shading)
    run = p.add_run('🔒  ' + text)
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(180, 35, 24)
    run.font.bold = True
    return p


def add_responsible_ai_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'E6F4E8')
    p._element.get_or_add_pPr().append(shading)
    run = p.add_run('✓  ' + text)
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(46, 125, 91)
    run.font.bold = True
    return p


def add_table_custom(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.allow_autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], '17365D')
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(36, 52, 71)
    return table


def add_cover(doc):
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Navy banner at top
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '17365D')
    p._element.get_or_add_pPr().append(shading)
    run = p.add_run('  AI CLASSROOM CHEATING DETECTION SYSTEM  ')
    run.font.color.rgb = RGBColor(255,255,255)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.name = 'Calibri'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run('Real-Time Exam Surveillance Using\nComputer Vision and Behavioral Analysis')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(23, 54, 93)
    run.font.name = 'Calibri'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run('Complete Project Documentation')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(15, 107, 120)
    run.font.name = 'Calibri'
    run.font.italic = True

    # Metadata box
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    table.allow_autofit = False
    hdr = table.rows[0].cells
    hdr[0].text = 'Project Identity'
    hdr[1].text = 'Document Control'
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True; run.font.name = 'Calibri'; run.font.size = Pt(10); run.font.color.rgb = RGBColor(255,255,255)
        set_cell_shading(cell, '17365D')

    meta = [
        ('Author', 'Md Moklesar Rahman'),
        ('Program', 'Master’s in Computer Science and Engineering'),
        ('Institution', 'Jahangirnagar University'),
        ('Supervisor', 'Risala Tasin Khan, PhD'),
        ('Email', 'md.moklasarrahmanbappy@gmail.com'),
        ('Phone', '+8801965031371'),
        ('Repository', 'https://github.com/Md-Moklesar-Rahman-Bappy/AI-Classroom-Cheating-Detection-System.git'),
        ('Document Version', '1.0 (final documentation generation)'),
        ('Repository Commit', '30a6ba2'),
        ('Prepared', '2026-08-30'),
        ('Status', 'Not production-ready; release criteria incomplete'),
        ('Document Type', 'Complete system, operator, developer, and research documentation'),
    ]
    for k, v in meta:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'; run.font.size = Pt(9)

    doc.add_paragraph()
    add_responsible_ai_box(doc, 'This document describes an AI-assisted surveillance system. AI-generated alerts indicate observable events that require human review. An alert is not proof of academic misconduct. Final academic or disciplinary decisions remain with authorized human reviewers and the institution.')

    doc.add_page_break()


def main():
    doc = Document()
    add_cover(doc)

    # Header / Footer
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = 'AI Classroom Cheating Detection — Complete Documentation | v1.0'
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in hp.runs:
            run.font.size = Pt(7); run.font.color.rgb = RGBColor(128,128,128); run.font.name = 'Calibri'
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = 'Commit 30a6ba2 — Not production-ready — Phase 10 Complete — 2026-08-30'
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in fp.runs:
            run.font.size = Pt(7); run.font.color.rgb = RGBColor(128,128,128); run.font.name = 'Calibri'

    # TOC placeholder
    add_heading_custom(doc, 'Table of Contents', 1)
    add_para(doc, 'Updated automatically in Microsoft Word (References → Table of Contents → Update Field). Sections include: Executive Summary, Project Background, Feature Catalog, System Architecture, CV/Event Logic, Database/API, Installation/Operation, Testing/Benchmarking, Security/Privacy, Appendices.')

    # PART I
    add_heading_custom(doc, 'Part I — Project Understanding', 1)
    add_heading_custom(doc, '1. Executive Summary', 2)
    add_para(doc, 'The AI Classroom Cheating Detection System is an AI-assisted examination-surveillance platform developed for research and institutional testing. It analyzes recorded or live video to detect observable events—person presence, mobile-phone visibility, orientation patterns, and temporal behavior changes—then produces evidence-containing reports for authorized human reviewers.')
    add_responsible_ai_box(doc, 'AI-generated alerts indicate observable events that require human review. An alert is not proof of academic misconduct.')

    add_heading_custom(doc, '2. Project Background', 2)
    add_para(doc, 'Developed as a Master’s project at Jahangirnagar University under supervisor Risala Tasin Khan, PhD. The repository is open-source (AGPL-3.0) and contains both synthetic/non-identifiable benchmark assets and infrastructure for approved staged recordings.')

    add_heading_custom(doc, '3. Problem Statement', 2)
    add_para(doc, 'Classroom examination integrity relies on human invigilator observation, which has coverage and consistency limits. The system provides AI-assisted observation to support—not replace—human review.')

    add_heading_custom(doc, '4. Motivation', 2)
    add_para(doc, 'Institutional need for consistent, documented, reviewable observation of examination sessions, with privacy controls, audit trails, and human-in-the-loop decisions.')

    add_heading_custom(doc, '5. Project Aim', 2)
    add_para(doc, 'Build a complete pipeline: camera/video input → frame processing → YOLO detection → anonymous tracking → orientation estimation → temporal rules → event generation → evidence management → dashboard review → audit log.')

    add_heading_custom(doc, '6. Specific Objectives', 2)
    objectives = [
        'Implement recorded-video analysis workflow (verified).',
        'Implement live-monitoring adapter (partially verified; webcam/test stream confirmed; RTSP/EZVIZ unverified).',
        'Implement anonymous tracking without identity recognition (verified).',
        'Implement geometric orientation method (experimental; verified in code; not validated for real-world accuracy).',
        'Implement temporal event rules for behavior patterns (partially implemented; synthetic evaluation only).',
        'Provide evidence-based human review (implemented).',
        'Maintain audit and retention controls (implemented).',
    ]
    for o in objectives:
        p = doc.add_paragraph(o, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.name = 'Calibri'; run.font.size = Pt(10)

    add_heading_custom(doc, '7. Research Questions', 2)
    add_para(doc, 'RQ1: Can synthetic/non-identifiable test material provide reproducible baseline metrics for object detection? (Answered: benchmark exists; real-data blocked.)')
    add_para(doc, 'RQ2: Does anonymous tracking maintain session-level consistency without identity recognition? (Answered: yes within session; identity switch possible at long occlusion.)')
    add_para(doc, 'RQ3: Can temporal rules distinguish sustained behavior from transient detection? (Answered: partially; requires more sequential data.)')

    add_heading_custom(doc, '8. Project Scope', 2)
    add_heading_custom(doc, '8.1 Included Scope', 3)
    add_para(doc, 'Recorded video upload, validation, storage, analysis, evidence generation, review workflow, audit logs, dashboard CRUD for video assets and analysis jobs, model version management, metrics/report generation, dataset governance for synthetic/non-identifiable data.')
    add_heading_custom(doc, '8.2 MVP Scope', 3)
    add_para(doc, 'Single-source recorded video → YOLO detection → events → evidence → review → report.')
    add_heading_custom(doc, '8.3 Experimental Scope', 3)
    add_para(doc, 'Live webcam/test stream monitoring; orientation estimation; temporal event engine; multi-source correlation (not fully validated).')
    add_heading_custom(doc, '8.4 Excluded Scope', 3)
    add_para(doc, 'Real participant recordings (blocked until consent verified); automated disciplinary decisions; facial recognition; biometric templates; public release of identifiable data.')
    add_heading_custom(doc, '8.5 Planned Future Scope', 3)
    add_para(doc, 'Verified EZVIZ integration (currently unverified); GPU benchmark; approved real-world dataset evaluation; improved temporal models with sufficient sequential data.')

    add_heading_custom(doc, '9. Expected Contribution', 2)
    add_para(doc, 'A reproducible, auditable, privacy-preserving AI-assisted examination-surveillance pipeline with documented governance, synthetic/non-identifiable evaluation, and explicit limitation statements.')

    add_heading_custom(doc, '10. Intended Users', 2)
    users = [
        ('System Administrator', 'Full access; configuration; model management; retention.'),
        ('Exam Administrator', 'Session/room management; review assignment.'),
        ('Invigilator', 'Live monitoring; evidence review.'),
        ('Reviewer', 'Event review; final decision; notes.'),
        ('Auditor', 'Audit logs; retention actions.'),
    ]
    for role, desc in users:
        add_para(doc, f'{role}: {desc}', bold=True)

    add_heading_custom(doc, '11. What the System Can Do', 2)
    add_para(doc, 'Detect visible persons and phones; track anonymously; estimate geometric orientation; apply temporal rules; generate evidence; support human review; log audits; manage video assets and jobs.')

    add_heading_custom(doc, '12. What the System Cannot Do', 2)
    add_para(doc, 'It does not automatically determine academic misconduct. It does not recognize identity. It does not guarantee real-world accuracy from synthetic benchmarks. It does not support unverified EZVIZ integration. It does not replace human review.')

    add_heading_custom(doc, '13. Important Limitations', 2)
    add_warning_box(doc, 'Real-data evaluation is BLOCKED (DATA_PLAN.md approved only for staged recordings with consent; synthetic/non-identifiable test data used currently).')
    add_warning_box(doc, 'Model accuracy on real classroom footage is NOT verified. Synthetic benchmark results do not generalize to all real conditions.')
    add_warning_box(doc, 'B4 Possible Seat Departure detection is partial/experimental unless fully validated.')
    add_warning_box(doc, 'Geometric orientation is a lightweight heuristic, not facial recognition or intent detection.')

    add_heading_custom(doc, '14. Responsible-AI Notice', 2)
    add_responsible_ai_box(doc, 'AI-generated alerts indicate observable events requiring human review. An alert is not proof of academic misconduct. Final academic or disciplinary decisions remain with authorized human reviewers and the institution.')

    # PART II FEATURE CATALOG
    add_heading_custom(doc, 'Part II — Feature Catalog', 1)
    add_heading_custom(doc, '22. Authentication', 2)
    add_para(doc, 'Implemented: Laravel Breeze-based auth with login, logout, password reset. Email verification available. Session encrypted (SESSION_ENCRYPT should be true in production; currently false in .env.local). Login rate limiting via Laravel throttle middleware.')
    add_security_box(doc, '.env local uses APP_DEBUG=true. Production must set APP_DEBUG=false.')

    add_heading_custom(doc, '23. Dashboard Overview', 2)
    add_para(doc, 'KPI cards, system status, recent activities, quick actions. All values derived from database queries (verified in controllers).')

    add_heading_custom(doc, '24. Exam Rooms', 2)
    add_para(doc, 'Created and managed via dashboard. Associated with sessions.')

    add_heading_custom(doc, '25. Exam Sessions', 2)
    add_para(doc, 'Created with name, status, created_by. Used by video assets and analysis jobs.')

    add_heading_custom(doc, '26. Camera Sources', 2)
    add_para(doc, 'Configured with type (webcam, test, RTSP), URL/credentials, preview settings.')

    add_heading_custom(doc, '27. Video Assets', 2)
    features = [
        'Upload with MIME validation (video/mp4, video/avi, video/quicktime, video/x-matroska).',
        'Private storage (stored_filename UUID-based).',
        'Status badges (pending/valid/invalid).',
        'Linked jobs count.',
        'Edit (session, filename, status).',
        'Soft delete with SweetAlert2 confirmation.',
        'Linked-job deletion protection.',
        'Audit entries for upload/update/delete/blocked.',
        'Re-upload for missing files.',
    ]
    for f in features:
        add_para(doc, '• ' + f, indent=False)

    add_heading_custom(doc, '28. Analysis Jobs', 2)
    add_para(doc, 'Create from session + video asset + model version + config. Job states: pending, queued, processing, completed, failed, cancelled. Status-based actions (view/edit/delete) configured in policy. Remote job ID and correlation ID tracked. Audit logs for create/approve/review/delete/retry.')

    add_heading_custom(doc, '29. Live Monitoring', 2)
    add_para(doc, 'Supported sources: webcam (verified), test stream (verified), RTSP adapter (implemented, unverified for EZVIZ), single-source limit (verified). States: running, paused, disconnected, degraded, reconnecting. Preview via MJPEG. Health endpoint verified.')
    add_warning_box(doc, 'EZVIZ CP1 Lite RTSP/ONVIF integration remains unverified.')

    add_heading_custom(doc, '30. Detection Events', 2)
    add_para(doc, 'Events generated by temporal rule engine: B1-B4 temporal events; D1/D2 detection outputs. Evidence snapshots include track number, event label, timestamp, and visual state. Human review required.')

    add_heading_custom(doc, '31. Evidence', 2)
    add_para(doc, 'Protected directory; not in Git; accessible to authorized reviewers only.')

    add_heading_custom(doc, '32. Human Reviews', 2)
    add_para(doc, 'Reviewer selects decision (confirmed suspicious, dismissed as normal, needs further review). Notes added. Audit log records decision with reviewer identifier.')

    add_heading_custom(doc, '33. Model Versions', 2)
    add_para(doc, 'Registered with checksum (yolo11n.pt: 0ebbc80d4... verified in benchmark_manifest). Weight file excluded from Git.')

    add_heading_custom(doc, '34. Metrics', 2)
    add_para(doc, 'Actual benchmark results in research/results/benchmark_results.json. No fabricated values.')

    add_heading_custom(doc, '35. Reports', 2)
    add_para(doc, 'PDF-style reports generated from job results and review decisions.')

    add_heading_custom(doc, '36. Audit Logs', 2)
    add_para(doc, 'Every job action (create, approve, review, delete, retry) logged with user, timestamp, action, result.')

    add_heading_custom(doc, '37. Users', 2)
    add_para(doc, 'Custom role system (system_admin, exam_admin, invigilator, reviewer, auditor) via role_user pivot.')

    add_heading_custom(doc, '38. Roles and Permissions', 2)
    add_para(doc, 'Policies registered (VideoAssetPolicy, AnalysisJobPolicy). Middleware enforces access.')

    add_heading_custom(doc, '39. Settings', 2)
    add_para(doc, 'AI service base URL, retry settings, timeout, model version defaults.')

    add_heading_custom(doc, '40. Help and Responsible-Use Notice', 2)
    add_responsible_ai_box(doc, 'This system is AI-assisted, not autonomous. Human review required.')

    # PART III ARCHITECTURE
    add_heading_custom(doc, 'Part III — System Architecture', 1)
    add_heading_custom(doc, '41. High-Level Architecture', 2)
    add_code_block(doc, 'Camera / Recorded Video → Adapter → Frame Scheduler → Preprocessing → YOLO Detector → Anonymous Tracker → Orientation Method → Temporal Rule Engine → Renderer / Evidence → Metrics / Report → API → Laravel Dashboard → Human Reviewer')

    add_heading_custom(doc, '42. Component Architecture', 2)
    add_para(doc, 'Components: input adapters (RecordedVideoInput, WebcamInput, RtspStreamInput, TestVideoInput), output adapters (EvidenceManager, MetricsExporter, ReportGenerator), detection engine (YOLO11 Nano), tracker (SimpleCentroidTracker), orientation (geometric-v1), temporal rules (ObservationWindow + RuleEngine).')

    add_heading_custom(doc, '43. AI-Service Architecture', 2)
    add_para(doc, 'FastAPI app; modules: main, api/jobs, api/live, api/health, api/version, api/events, api/metrics, config/settings. Uses uvicorn; must start from ai-service root.')

    add_heading_custom(doc, '44. Laravel Dashboard Architecture', 2)
    add_para(doc, 'MVC with Blade views; roles via custom pivot; policies registered; audit via AuditHelper; soft deletes on VideoAsset and AnalysisJob.')

    add_heading_custom(doc, '45. Shared Detection Engine', 2)
    add_para(doc, 'Python modules in ai-service/app/ detect person and mobile-phone; feed tracker; compute orientation; evaluate temporal rules.')

    add_heading_custom(doc, '46. Input Adapters (verified)', 2)
    add_table_custom(doc,
        ['Adapter', 'Type', 'Status'],
        [
            ['RecordedVideoInput', 'Recorded file', 'Implemented'],
            ['WebcamInput', 'Local webcam', 'Implemented'],
            ['RtspStreamInput', 'RTSP stream', 'Implemented; unverified for EZVIZ'],
            ['TestVideoInput', 'Synthetic test', 'Implemented'],
        ])

    add_heading_custom(doc, '47. Output Adapters', 2)
    add_para(doc, 'EvidenceManager (snapshots), MetricsExporter (benchmark), ReportGenerator (PDF-style), AuditLogger (database).')

    add_heading_custom(doc, '48. Recorded-Video Workflow', 2)
    add_para(doc, 'Upload → validation → storage → analysis job creation → remote job → poll → events → evidence → review → metrics.')

    add_heading_custom(doc, '49. Live-Surveillance Workflow', 2)
    add_para(doc, 'Start source → frame scheduler → detection → tracker → orientation → temporal rules → event alert → evidence → preview → stop.')

    add_heading_custom(doc, '50. Cross-Service Multipart-Transfer Workflow', 2)
    add_para(doc, 'Laravel asset → stream open → multipart → FastAPI validation → temporary input → remote job → remote ID → polling → event sync → evidence transfer → metrics sync → cleanup.')

    add_heading_custom(doc, '51. Queue Architecture', 2)
    add_para(doc, 'Laravel queue (database driver); worker must remain running. Job retries configurable.')

    add_heading_custom(doc, '52. Job Lifecycle and State Machine', 2)
    add_para(doc, 'States: pending → queued → processing → completed / failed / cancelled. Retry and cancel supported.')

    add_heading_custom(doc, '53. Event Lifecycle', 2)
    add_para(doc, 'Detection → observation window → supporting observation → event end → evidence → review.')

    add_heading_custom(doc, '54. Evidence Lifecycle', 2)
    add_para(doc, 'Created on event confirmation; protected; retained per policy; deleted secure if required.')

    add_heading_custom(doc, '55. Human-Review Lifecycle', 2)
    add_para(doc, 'Pending review → reviewer selects → notes added → audit log → evidence updated → report updated.')

    add_heading_custom(doc, '56. Audit Lifecycle', 2)
    add_para(doc, 'Every action logged with user, action, result, timestamp.')

    add_heading_custom(doc, '57. Error Handling and Recovery', 2)
    add_para(doc, 'Service unavailable: retry; timeout: configurable; invalid stream: drop; unauthorized access: 403; missing file: re-upload or block.')

    add_heading_custom(doc, '58. Correlation IDs and Traceability', 2)
    add_para(doc, 'Correlation ID (user-set) and remote job ID (service-returned) allow trace from dashboard to AI output.')

    # PART IV CV / EVENT
    add_heading_custom(doc, 'Part IV — Computer Vision and Event Logic', 1)
    add_heading_custom(doc, '59. Model Baseline', 2)
    add_para(doc, 'YOLO11 Nano (yolo11n.pt). Classes: person (COCO 0), cell phone (COCO 66). Checksum verified: 0ebbc80d4a7680d14987a577cd21342b65ecfd94632bd9a8da63ae4617644ee1 (benchmark_manifest.json). Source: Ultralytics. License: AGPL-3.0. CPU inference only (current environment). Weight file not committed (excluded by .gitignore).')
    add_warning_box(doc, 'No real-world accuracy study performed; synthetic benchmark only.')

    add_heading_custom(doc, '60. Person Detection', 2)
    add_para(doc, 'Bounding box around visible torso/head. Not a suspicious event by itself.')

    add_heading_custom(doc, '61. Mobile-Phone Detection', 2)
    add_para(doc, 'Visible handheld device. Not proof of usage.')

    add_heading_custom(doc, '62. Anonymous Tracking', 2)
    add_para(doc, 'SimpleCentroidTracker; temporary IDs; distance threshold; missed-frame tolerance. No embeddings. Session scope. Identity-switch possible at long occlusion.')

    add_heading_custom(doc, '63. Orientation Method', 2)
    add_para(doc, 'Geometric-v1: uses relative positions of tracked points to derive direction states. Lightweight heuristic. Not facial recognition; not intent detection.')

    add_heading_custom(doc, '64. Temporal Rule Engine', 2)
    add_para(doc, 'Observation window; supporting-observation count; minimum duration; cooldown; duplicate suppression.')

    add_heading_custom(doc, '65. Event Taxonomy', 2)
    add_table_custom(doc, ['Category', 'Code', 'Meaning', 'Status'],
        [
            ['Detection', 'D1', 'Person detected', 'Implemented'],
            ['Detection', 'D2', 'Mobile phone detected', 'Implemented'],
            ['Behavior', 'B1', 'Repeated looking left', 'Partial / experimental'],
            ['Behavior', 'B2', 'Repeated looking right', 'Partial / experimental'],
            ['Behavior', 'B3', 'Looking backward', 'Partial / experimental'],
            ['Behavior', 'B4', 'Possible seat departure', 'Partial / experimental'],
            ['State', 'S1', 'Normal', 'Implemented'],
            ['State', 'S2', 'Insufficient evidence', 'Implemented'],
        ])

    add_heading_custom(doc, '66. Bounding Boxes and Visual States', 2)
    add_para(doc, 'Green (normal), Amber (uncertain / needs review), Red (suspicious event observed — requires human decision), Blue (evidence selected), Gray (inactive / unavailable).')

    add_heading_custom(doc, '67. B4 Leaving-Seat Limitations', 2)
    add_warning_box(doc, 'B4 Possible Seat Departure is partial/experimental. Not fully validated for real-world conditions. Requires temporal evidence and human confirmation.')

    # PART V DATABASE / API
    add_heading_custom(doc, 'Part V — Database and API', 1)
    add_heading_custom(doc, '68. Database Overview', 2)
    add_para(doc, 'MySQL (configured via DB_CONNECTION=mysql). Migration-backed schema. Soft deletes on VideoAsset and AnalysisJob.')

    add_heading_custom(doc, '69. ER Overview', 2)
    add_para(doc, 'Key entities: users → roles → exam_sessions → video_assets → analysis_jobs → detection_events → event_evidence → review_decisions → audit_logs. VideoAsset links to AnalysisJob via video_asset_id.')

    add_heading_custom(doc, '70. Table Dictionary (verified)', 2)
    tables = [
        ['users', 'Authentication', 'id, name, email, password, role_id, created_at'],
        ['exam_sessions', 'Session metadata', 'id, name, status, exam_room_id, created_by, created_at'],
        ['video_assets', 'Uploaded videos', 'id, exam_session_id, stored_filename, mime_type, size_bytes, validation_status, deleted_at'],
        ['analysis_jobs', 'Processing jobs', 'id, exam_session_id, video_asset_id, source_type, model_version_id, status, remote_job_id, correlation_id, deleted_at'],
        ['model_versions', 'Model registry', 'id, name, version, checksum_sha256, weight_filename'],
        ['detection_events', 'Events', 'id, analysis_job_id, event_type, start_frame, end_frame, confidence'],
        ['event_evidence', 'Evidence', 'id, event_id, snapshot_path, review_decision_id'],
        ['review_decisions', 'Review', 'id, evidence_id, reviewer_code, decision, notes'],
        ['audit_logs', 'Audit', 'id, user_id, action, model, record_id, result'],
    ]
    add_table_custom(doc, ['Table', 'Purpose', 'Key Columns'], tables)

    add_heading_custom(doc, '71. Model Relationships', 2)
    add_para(doc, 'AnalysisJob belongsTo VideoAsset (video_asset_id). VideoAsset hasMany AnalysisJob (analysisJobs). DetectionEvent belongsTo AnalysisJob.')

    add_heading_custom(doc, '72. Soft-Delete Strategy', 2)
    add_para(doc, 'VideoAsset and AnalysisJob both use SoftDeletes. Deleted records excluded from normal queries; recoverable via restore(). Physical file deletion is separate from soft delete.')

    add_heading_custom(doc, '73. API Overview', 2)
    add_para(doc, 'FastAPI endpoints: /health, /version, /jobs (create/status/cancel/retry), /events, /metrics, /live/start/stop/health/preview.')

    add_heading_custom(doc, '74. Authentication and Headers', 2)
    add_para(doc, 'API token via env (AI_SERVICE_TOKEN). Not exposed in Git. HTTPS recommended in production.')

    add_heading_custom(doc, '75. Endpoint Reference (verified from routes)', 2)
    endpoints = [
        ['GET', '/health', 'Service health', 'None', '200 / 503'],
        ['GET', '/version', 'Model/version info', 'None', '200'],
        ['POST', '/jobs', 'Create analysis job', 'Token', '201 / 400 / 401'],
        ['GET', '/jobs/{id}', 'Job status', 'Token', '200 / 404'],
        ['POST', '/jobs/{id}/cancel', 'Cancel', 'Token', '200'],
        ['POST', '/jobs/{id}/retry', 'Retry', 'Token', '200'],
        ['GET', '/events', 'Event list', 'Token', '200'],
        ['GET', '/metrics', 'Performance metrics', 'Token', '200'],
        ['POST', '/live/start', 'Start live', 'Token', '200'],
        ['POST', '/live/stop', 'Stop live', 'Token', '200'],
    ]
    add_table_custom(doc, ['Method', 'Path', 'Purpose', 'Auth', 'Status'], endpoints)

    add_heading_custom(doc, '76. Multipart Upload Specification', 2)
    add_para(doc, 'multipart/form-data with file field; validated by FastAPI; temporary storage during processing.')

    add_heading_custom(doc, '77. Idempotency', 2)
    add_para(doc, 'Retry attempts configurable; correlation ID allows duplicate detection.')

    add_heading_custom(doc, '78. Error Response Format', 2)
    add_para(doc, 'JSON with error code, message, and details.')

    add_heading_custom(doc, '79. Secret Redaction', 2)
    add_security_box(doc, '.env excluded from Git; token via env only; config default is placeholder.')

    # PART VI INSTALLATION / OPERATION
    add_heading_custom(doc, 'Part VI — Installation and Configuration', 1)
    add_heading_custom(doc, '80. Prerequisites', 2)
    add_para(doc, 'Windows 10/11 (verified environment); PHP 8+; Composer; MySQL/MariaDB; Node.js/npm; Python 3.14; OpenCV; Ultralytics; PyTorch CPU.')

    add_heading_custom(doc, '81. Repository Cloning', 2)
    add_code_block(doc, 'git clone https://github.com/Md-Moklesar-Rahman-Bappy/AI-Classroom-Cheating-Detection-System.git')

    add_heading_custom(doc, '82. Directory Structure', 2)
    add_para(doc, 'Root contains ai-service/, dashboard/, docs/, research/, scripts/. Data and weights outside Git.')

    add_heading_custom(doc, '83. Python Environment Setup', 2)
    add_code_block(doc, r"cd ai-service\npython -m venv .venv\nsource .venv/bin/activate  # Git Bash\n.venv\Scripts\activate    # PowerShell\npip install -r requirements-dev.txt")
    add_important_box(doc, 'Use python -m uvicorn, not bare uvicorn, to avoid PATH issues observed during development.')

    add_heading_custom(doc, '84. AI-Service Environment', 2)
    add_para(doc, 'Copy .env.example to .env; set AI_SERVICE_TOKEN (not default placeholder in production); configure base_url, timeouts.')

    add_heading_custom(doc, '85. Starting FastAPI', 2)
    add_code_block(doc, 'python -m uvicorn app.main:app --host 127.0.0.1 --port 8001')
    add_important_box(doc, 'Start from ai-service root. Previous ModuleNotFoundError caused by wrong working directory.')

    add_heading_custom(doc, '86. AI Health Validation', 2)
    add_para(doc, 'GET http://127.0.0.1:8001/health; verify version endpoint.')

    add_heading_custom(doc, '87. Laravel Setup', 2)
    add_code_block(doc, 'cd dashboard\ncomposer install\ncp .env.example .env\nphp artisan key:generate\n# configure DB in .env\nphp artisan migrate --force\nphp artisan db:seed\nnpm install\nnpm run build')

    add_heading_custom(doc, '88. Database Configuration', 2)
    add_para(doc, 'MySQL connection; database ai_classroom; user root (development only); password empty in .env.local.')

    add_heading_custom(doc, '89. Role / Default Model Seeding', 2)
    add_para(doc, 'Seeders create roles (system_admin, exam_admin, invigilator, reviewer, auditor) and default model version.')

    add_heading_custom(doc, '90. Starting Laravel', 2)
    add_code_block(doc, 'php artisan serve')

    add_heading_custom(doc, '91. Starting Queue Worker', 2)
    add_code_block(doc, 'php artisan queue:work --queue=default --tries=1 --timeout=900 -v')
    add_warning_box(doc, 'Worker must remain running. Jobs stay pending without it. Restart after code/config changes: php artisan queue:restart.')

    add_heading_custom(doc, '92. Full Three-Terminal Startup', 2)
    add_table_custom(doc, ['Terminal', 'Command', 'Purpose'],
        [['T1', 'python -m uvicorn ...', 'AI service'], ['T2', 'php artisan serve', 'Dashboard'], ['T3', 'php artisan queue:work ...', 'Worker']])

    add_heading_custom(doc, '93. First-Run Verification', 2)
    add_para(doc, 'Check login, video asset upload, analysis job creation, job processing, event generation, evidence, review.')

    add_heading_custom(doc, '94. Environment Variables', 2)
    add_table_custom(doc, ['Key', 'Purpose', 'Required', 'Security'],
        [
            ['APP_KEY', 'Encryption', 'Yes', 'Secret; do not commit'],
            ['DB_PASSWORD', 'DB access', 'Yes', 'Secret'],
            ['AI_SERVICE_TOKEN', 'API auth', 'Yes', 'Secret; placeholder in config'],
            ['APP_DEBUG', 'Debug mode', 'No (set false in prod)', 'Not secret'],
            ['AI_SERVICE_BASE_URL', 'Service URL', 'No', 'Not secret'],
            ['SESSION_ENCRYPT', 'Session encrypt', 'No', 'Should be true in prod'],
        ])

    # PART VII OPERATOR MANUAL
    add_heading_custom(doc, 'Part VII — Step-by-Step Operator Manual', 1)
    add_heading_custom(doc, '95. Signing In', 2)
    add_para(doc, 'Use login page with valid account assigned to a role.')

    add_heading_custom(doc, '96. Understanding the Sidebar', 2)
    add_para(doc, 'Navigation: Dashboard, Exam Rooms, Exam Sessions, Cameras, Video Assets, Analysis Jobs, Live Monitoring, Evidence, Reviews, Reports, Audit Logs.')

    add_heading_custom(doc, '97. Signing Out', 2)
    add_para(doc, 'Use top-right user menu.')

    add_heading_custom(doc, '98. Creating an Exam Room', 2)
    add_para(doc, 'Fill name and description; save.')

    add_heading_custom(doc, '99. Creating an Exam Session', 2)
    add_para(doc, 'Select exam room; enter name; choose status; save.')

    add_heading_custom(doc, '100. Uploading a Recorded Video', 2)
    add_para(doc, 'Supported formats: mp4, avi, mov, mkv. Validation checks MIME and size (max ~512 MB). Storage is private (UUID filename). Status set to pending/valid/invalid based on validation.')

    add_heading_custom(doc, '101. Viewing Video Assets', 2)
    add_para(doc, 'Index table shows SL, original, stored (truncated), MIME, size, created date, status badge, linked jobs count, actions.')

    add_heading_custom(doc, '102. Editing a Video Asset', 2)
    add_para(doc, 'Click Edit to change exam session, original filename, validation status.')

    add_heading_custom(doc, '103. Deleting a Video Asset', 2)
    add_para(doc, 'Click Delete → SweetAlert2 confirmation → soft delete (recoverable). Blocked if analysis jobs linked (shows count). Audit log entry created.')

    add_heading_custom(doc, '104. Starting Analysis from Asset', 2)
    add_para(doc, 'Click Analyze from asset page → creates job with pre-selected video asset.')

    add_heading_custom(doc, '105. Creating Analysis Job Manually', 2)
    add_para(doc, 'Select session → source type → video asset → model version → config (interval, mode, resolution). Submit.')

    add_heading_custom(doc, '106. Selecting Correct Video Asset', 2)
    add_para(doc, 'Only valid assets shown; missing file assets require re-upload.')

    add_heading_custom(doc, '107. Selecting the Model', 2)
    add_para(doc, 'Only registered model versions appear (e.g., yolo11n).')

    add_heading_custom(doc, '108. Selecting Configuration', 2)
    add_para(doc, 'Interval (every frame / every 3rd / every 5th); mode (recorded/live); resolution (640x360 / 480x270).')

    add_heading_custom(doc, '109. Understanding Job Statuses', 2)
    add_table_custom(doc, ['Status', 'Meaning', 'Allowed Actions'],
        [['Pending', 'Queued for worker', 'Edit / Cancel'], ['Queued', 'In queue', 'View'], ['Processing', 'Running', 'View'], ['Completed', 'Finished', 'View / Report / Delete'], ['Failed', 'Error', 'Retry / View'], ['Cancelled', 'Stopped', 'Retry']])

    add_heading_custom(doc, '110. Running the Queue Worker', 2)
    add_para(doc, 'Keep terminal open. Restart after code changes.')

    add_heading_custom(doc, '111. Monitoring Progress', 2)
    add_para(doc, 'View job list; refresh for updated status and remote job ID.')

    add_heading_custom(doc, '112. Remote Job ID', 2)
    add_para(doc, 'Returned by FastAPI; used for correlation.')

    add_heading_custom(doc, '113. Correlation ID', 2)
    add_para(doc, 'User-set or auto-generated; links dashboard to AI output.')

    add_heading_custom(doc, '114. Cancelling a Job', 2)
    add_para(doc, 'Available for pending/queued/processing.')

    add_heading_custom(doc, '115. Retrying a Failed Job', 2)
    add_para(doc, 'Retry action creates new attempt with same config.')

    add_heading_custom(doc, '116. Editing a Pending Job', 2)
    add_para(doc, 'Edit session, asset, model, config.')

    add_heading_custom(doc, '117. Deleting an Analysis Job', 2)
    add_para(doc, 'Soft delete; audit log.')

    add_heading_custom(doc, '118. Viewing Annotated Result', 2)
    add_para(doc, 'Click completed job → view output path → review events/evidence.')

    add_heading_custom(doc, '119. Understanding Events', 2)
    add_para(doc, 'Events shown with start/end frames, confidence, types.')

    add_heading_custom(doc, '120. Viewing Evidence', 2)
    add_para(doc, 'Evidence snapshot with event label, review status.')

    add_heading_custom(doc, '121. Submitting Human Review', 2)
    add_para(doc, 'Select decision; add notes; submit.')

    add_heading_custom(doc, '122. Adding Reviewer Notes', 2)
    add_para(doc, 'Free-text notes stored with review.')

    add_heading_custom(doc, '123. Generating a Report', 2)
    add_para(doc, 'Report includes job info, events, evidence, review, metrics.')

    add_heading_custom(doc, '124. Reading Metrics', 2)
    add_para(doc, 'Actual benchmark results from research/results/benchmark_results.json.')

    add_heading_custom(doc, '125. Managing Model Versions', 2)
    add_para(doc, 'Register with checksum; weight file not in Git.')

    add_heading_custom(doc, '126. Managing Users and Roles', 2)
    add_para(doc, 'Assign roles through admin interface.')

    add_heading_custom(doc, '127. Viewing Audit Logs', 2)
    add_para(doc, 'Search by user, action, model, result.')

    add_heading_custom(doc, '128. Operating Live Mode (Webcam)', 2)
    add_para(doc, 'Select webcam source → Start → Monitor → Stop.')

    add_heading_custom(doc, '129. Operating Live Mode (Test Stream)', 2)
    add_para(doc, 'Select test source → Start.')

    add_heading_custom(doc, '130. Attempting RTSP Source', 2)
    add_warning_box(doc, 'EZVIZ CP1 Lite RTSP/ONVIF support remains unverified.')

    add_heading_custom(doc, '131. Stopping Live Monitoring', 2)
    add_para(doc, 'Click Stop; verify state returns to normal.')

    add_heading_custom(doc, '132. Responding to Degraded State', 2)
    add_para(doc, 'Check camera; reconnect; review logs.')

    add_heading_custom(doc, '133. Safe Shutdown', 2)
    add_para(doc, 'Stop live; complete or cancel jobs; stop queue worker; stop FastAPI; close terminal.')

    # PART VIII ADMIN MANUAL
    add_heading_custom(doc, 'Part VIII — Administrator Manual', 1)
    add_heading_custom(doc, '134. Role Matrix', 2)
    add_table_custom(doc, ['Role', 'Can Edit/Del Video', 'Can Manage Jobs', 'Can Review', 'Can Audit', 'Can Manage Users'],
        [['system_admin', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
         ['exam_admin', 'Yes', 'Yes', 'Yes', 'Yes', 'No'],
         ['invigilator', 'No', 'View', 'No', 'No', 'No'],
         ['reviewer', 'No', 'View', 'Yes', 'No', 'No'],
         ['auditor', 'No', 'View', 'No', 'Yes', 'No']])

    add_heading_custom(doc, '135. Assigning Roles', 2)
    add_para(doc, 'Custom role_user pivot; no Spatie package used.')

    add_heading_custom(doc, '136. Restoring Seeded Roles', 2)
    add_para(doc, 'Run seeders or insert into database.')

    add_heading_custom(doc, '137. Restoring Default Model', 2)
    add_para(doc, 'Re-register yolo11n version with correct checksum.')

    add_heading_custom(doc, '138. Managing Storage', 2)
    add_para(doc, 'Private disk; backup manually; exclude from Git.')

    add_heading_custom(doc, '139. Managing Queue Workers', 2)
    add_para(doc, 'Keep running; restart after changes.')

    add_heading_custom(doc, '140. Managing Failed Jobs', 2)
    add_para(doc, 'View failed-jobs table; retry or delete.')

    add_heading_custom(doc, '141. Clearing Caches', 2)
    add_code_block(doc, 'php artisan optimize:clear')

    add_heading_custom(doc, '142. Restarting Services', 2)
    add_para(doc, 'Restart FastAPI, Laravel server, and queue worker.')

    add_heading_custom(doc, '143. Log Management', 2)
    add_para(doc, 'Laravel logs in storage/logs/; AI service logs to stdout/stderr.')

    add_heading_custom(doc, '144. Backup Considerations', 2)
    add_para(doc, 'Backup database (sql dump), storage directory (excluding raw videos if required by policy), and model weights separately.')

    add_heading_custom(doc, '145. Evidence Retention and Deletion', 2)
    add_para(doc, 'Follow DATA_RETENTION_POLICY.md. Secure deletion required for early removal.')

    add_heading_custom(doc, '146. Missing-File Reconciliation', 2)
    add_para(doc, 'Video asset shows invalid when file missing; re-upload from backup or source.')

    add_heading_custom(doc, '147. Security Incident Response', 2)
    add_para(doc, 'Document in audit; review access; check retention actions; report to supervisor.')

    add_heading_custom(doc, '148. Production Configuration Checklist', 2)
    add_warning_box(doc, 'Not production-ready (release criteria incomplete). Before production: fix .env secrets, enable session encryption, disable debug, set HTTPS, configure CI/dependabot, complete user acceptance, verify GPU benchmark if needed.')

    # PART IX DEVELOPER / MAINTAINER
    add_heading_custom(doc, 'Part IX — Developer and Maintainer Guide', 1)
    add_heading_custom(doc, '149. Source-Code Organization', 2)
    add_para(doc, 'AI service: Python modules; Dashboard: Laravel MVC; Tests: PHPUnit + Python.')

    add_heading_custom(doc, '150. AI-Service Modules', 2)
    add_para(doc, 'main, api/jobs, api/live, api/health, api/version, api/events, api/metrics, config/settings.')

    add_heading_custom(doc, '151. Dashboard Modules', 2)
    add_para(doc, 'Controllers, Models (with SoftDeletes), Policies, Middleware, Views, Routes.')

    add_heading_custom(doc, '152. Configuration Architecture', 2)
    add_para(doc, '.env + config/*.php; AI config separate from Laravel.')

    add_heading_custom(doc, '153. Adding an Input Adapter', 2)
    add_para(doc, 'Implement adapter class; register in scheduler; add to allowed source types.')

    add_heading_custom(doc, '154. Adding an Output Adapter', 2)
    add_para(doc, 'Implement adapter; register in pipeline.')

    add_heading_custom(doc, '155. Adding a Detector Output', 2)
    add_para(doc, 'Add class to taxonomy; update annotation guide; update metrics if needed.')

    add_heading_custom(doc, '156. Adding a Temporal Rule', 2)
    add_para(doc, 'Add to rule engine; set observation window and duration.')

    add_heading_custom(doc, '157. Adding an Event Type', 2)
    add_para(doc, 'Update taxonomy, annotation guide, evidence manager.')

    add_heading_custom(doc, '158. Updating Event Taxonomy Safely', 2)
    add_para(doc, 'Keep backward compatibility; document version change.')

    add_heading_custom(doc, '159. Updating Database Schema', 2)
    add_code_block(doc, 'php artisan make:migration add_x_to_y_table\n# edit migration\nphp artisan migrate')

    add_heading_custom(doc, '160. Adding an API Endpoint', 2)
    add_para(doc, 'Add to FastAPI routes; document in API reference.')

    add_heading_custom(doc, '161. Adding a Dashboard Module', 2)
    add_para(doc, 'Controller + model + policy + view + route.')

    add_heading_custom(doc, '162. Adding Authorization Policy', 2)
    add_para(doc, 'Create Policy class; register in AppServiceProvider.')

    add_heading_custom(doc, '163. Adding Audit Events', 2)
    add_para(doc, 'Use AuditHelper::log().')

    add_heading_custom(doc, '164. Adding Tests', 2)
    add_para(doc, 'Laravel feature tests; Python tests for AI service.')

    add_heading_custom(doc, '165. Logging Guidelines', 2)
    add_para(doc, 'Use structured JSON where possible; avoid secrets.')

    add_heading_custom(doc, '166. Secret-Handling Guidelines', 2)
    add_security_box(doc, 'No secrets in source. Use .env. Default token is placeholder only.')

    add_heading_custom(doc, '167. Model-Version Management', 2)
    add_para(doc, 'Register with checksum; store weight outside Git.')

    add_heading_custom(doc, '168. Dependency Updates', 2)
    add_para(doc, 'Update requirements; test; document in notes.')

    add_heading_custom(doc, '169. Git and Commit Conventions', 2)
    add_para(doc, 'Use descriptive messages; include issue reference; do not commit secrets or weights.')

    add_heading_custom(doc, '170. Branching and Release Considerations', 2)
    add_para(doc, 'Release not tagged (authorization required). Use feature branches.')

    add_heading_custom(doc, '171. Coding Standards', 2)
    add_para(doc, 'Follow Laravel conventions; Python PEP8; document new features.')

    add_heading_custom(doc, '172. Safe Extension Checklist', 2)
    add_para(doc, 'Verify authorization, audit, soft delete, validation, tests.')

    # PART X TESTING / BENCHMARKING / EVALUATION
    add_heading_custom(doc, 'Part X — Testing, Benchmarking, and Evaluation', 1)
    add_heading_custom(doc, '173. Test Strategy', 2)
    add_para(doc, 'Laravel PHPUnit feature tests (video assets, analysis jobs). Python benchmark script (actual execution, not invented).')

    add_heading_custom(doc, '174. Running Python Tests', 2)
    add_para(doc, 'No dedicated Python test folder; benchmark script serves as validation.')

    add_heading_custom(doc, '175. Running Laravel Tests', 2)
    add_code_block(doc, 'php artisan test tests/Feature/VideoAssetFailureTest.php')

    add_heading_custom(doc, '176. Running Ruff', 2)
    add_para(doc, 'Not executed in session (recommended for future).')

    add_heading_custom(doc, '177. Running Black', 2)
    add_para(doc, 'Not executed in session.')

    add_heading_custom(doc, '178. Running Pint', 2)
    add_para(doc, 'Not executed in session.')

    add_heading_custom(doc, '179. Running Composer Validation', 2)
    add_para(doc, 'Not executed in session.')

    add_heading_custom(doc, '180. Running Production Asset Build', 2)
    add_para(doc, 'Not executed in session (npm run build recommended).')

    add_heading_custom(doc, '181. Static Analysis Status', 2)
    add_para(doc, 'Not fully executed (some LSP diagnostics observed; not quality results).')

    add_heading_custom(doc, '182. Test Fixtures', 2)
    add_para(doc, 'Synthetic fixtures in storage/ and benchmark_manifest.json.')

    add_heading_custom(doc, '183. Difference Between Tests and Real-World Accuracy', 2)
    add_warning_box(doc, 'Unit tests validate software behavior. They do not prove real-world model accuracy on classroom footage.')

    add_heading_custom(doc, '184. Benchmark Methodology (Phase 8)', 2)
    add_para(doc, 'Environment: CPU-only (Intel i5-14500, 8 GB RAM, no GPU). Synthetic fixture: 640x360, 10 fps, 90 frames, gray rectangle + white circle, no person, no PII. Resolutions: 640x360 and 480x270. Frame intervals: every 1, 3, 5 frames. Warm-up separate. Actual measured FPS, latency, CPU, memory from benchmark_manifest.json.')

    add_heading_custom(doc, '185. Low-Resource Profile', 2)
    add_para(doc, 'CPU inference; no GPU required for basic operation; performance lower than GPU-equipped systems.')

    add_heading_custom(doc, '186. Real Recorded-Job Runtime Example', 2)
    add_para(doc, 'Job 12 completed successfully (verified in project history). Exact metrics not preserved in repository files; not presented as accuracy study.')

    add_heading_custom(doc, '187. Dataset Governance', 2)
    add_para(doc, 'Documented; consent materials present; real-data blocked.')

    add_heading_custom(doc, '188. Consent', 2)
    add_para(doc, 'CONSENT_TEMPLATE.md present.')

    add_heading_custom(doc, '189. Data Collection', 2)
    add_para(doc, 'DATA_COLLECTION_PROTOCOL.md present.')

    add_heading_custom(doc, '190. Annotation Protocol', 2)
    add_para(doc, 'ANNOTATION_GUIDE.md present.')

    add_heading_custom(doc, '191. Dataset Manifest', 2)
    add_para(doc, 'research/manifests/MANIFEST.json (v0.1.0-synthetic).')

    add_heading_custom(doc, '192. Dataset Versioning', 2)
    add_para(doc, 'DATASET_VERSIONING.md present.')

    add_heading_custom(doc, '193. Split Policy', 2)
    add_para(doc, 'DATASET_SPLIT_POLICY.md present.')

    add_heading_custom(doc, '194. Leakage Prevention', 2)
    add_para(doc, 'check_split_leakage.py verifies.')

    add_heading_custom(doc, '195. Evaluation Scripts', 2)
    add_para(doc, 'evaluate_objects.py, evaluate_events.py, check_split_leakage.py, generate_confusion_matrix.py, check_data_validation.py.')

    add_heading_custom(doc, '196. Evaluation Metrics', 2)
    add_para(doc, 'Actual results in benchmark_results.json and sanitized_evaluation_result.json. No fabricated metrics.')

    add_heading_custom(doc, '197. Current Evaluation Status', 2)
    add_para(doc, 'Synthetic/non-identifiable evaluation active. Real participant evaluation BLOCKED (no consent artifacts verified for real recordings).')

    add_heading_custom(doc, '198. Reproducibility', 2)
    add_para(doc, 'REPRODUCIBILITY.md; seed and hardware recorded; commands documented.')

    # PART XI SECURITY / PRIVACY / ETHICS / LICENSING
    add_heading_custom(doc, 'Part XI — Security, Privacy, Ethics, and Licensing', 1)
    add_heading_custom(doc, '199. Security Architecture', 2)
    add_para(doc, 'Laravel auth; policies; middleware; soft deletes; audit logs; private storage; secret redaction; HTTPS required in production.')

    add_heading_custom(doc, '200. Authentication Controls', 2)
    add_para(doc, 'Password-based; session encrypted (recommended true in prod); rate limiting via throttle.')

    add_heading_custom(doc, '201. Authorization Controls', 2)
    add_para(doc, 'Role-based (system_admin, exam_admin, invigilator, reviewer, auditor). Policies enforce per-resource access.')

    add_heading_custom(doc, '202. Upload Security', 2)
    add_para(doc, 'MIME validation; size limit; private storage; UUID filenames.')

    add_heading_custom(doc, '203. API Security', 2)
    add_para(doc, 'Token auth; HTTPS recommended; no secrets in source.')

    add_heading_custom(doc, '204. Camera-Credential Protection', 2)
    add_para(doc, 'Credentials stored in environment/config only; not in Git.')

    add_heading_custom(doc, '205. Evidence Protection', 2)
    add_para(doc, 'Evidence directory excluded from Git; access controlled by authorization.')

    add_heading_custom(doc, '206. Secret Redaction', 2)
    add_security_box(doc, '.env excluded; no committed secrets found; placeholder token documented.')

    add_heading_custom(doc, '207. Audit Logging', 2)
    add_para(doc, 'Every action logged with user, action, model, result.')

    add_heading_custom(doc, '208. Data Minimization', 2)
    add_para(doc, 'No names in manifest; temporary track IDs only.')

    add_heading_custom(doc, '209. Retention', 2)
    add_para(doc, '90-day default; secure deletion; early deletion allowed.')

    add_heading_custom(doc, '210. Secure Deletion', 2)
    add_para(doc, 'Overwrite/unlink; verify missing; log action.')

    add_heading_custom(doc, '211. Threat Model Summary', 2)
    add_para(doc, 'Threat model documented (THREAT_MODEL.md). Risks: unauthorized access, data leakage, model misuse, evidence tampering.')

    add_heading_custom(doc, '212. Responsible-AI Controls', 2)
    add_responsible_ai_box(doc, 'Human review required; no automated disciplinary decision; disclaimer in documentation and UI.')

    add_heading_custom(doc, '213. Human-in-the-Loop Requirement', 2)
    add_para(doc, 'Every alert requires reviewer confirmation before any action.')

    add_heading_custom(doc, '214. Privacy Limitations', 2)
    add_para(doc, 'No real-participant data included; synthetic only; consent artifacts present for future approved collection.')

    add_heading_custom(doc, '215. Ethical Data-Collection Requirements', 2)
    add_para(doc, 'Institutional approval + informed consent + retention notice required before any real recording.')

    add_heading_custom(doc, '216. Third-Party Dependencies', 2)
    add_table_custom(doc, ['Package', 'Version', 'License', 'Usage', 'Notes'],
        [['ultralytics', '8.4.135', 'AGPL-3.0', 'YOLO model', 'Weight not redistributed'],
         ['torch', '2.13.0+cpu', 'BSD', 'Framework', 'CPU-only'],
         ['opencv', '5.0.0', 'BSD', 'Image processing', 'Verified'],
         ['fastapi', 'Verified', 'MIT', 'API', 'Verified'],
         ['laravel/framework', 'Verified', 'MIT', 'Dashboard', 'Verified'],
         ['mysql', 'Verified', 'GPL', 'Database', 'Verified']])

    add_heading_custom(doc, '217. Ultralytics AGPL-3.0', 2)
    add_para(doc, 'Package: ultralytics 8.4.135. License: AGPL-3.0. Usage: YOLO detection. Source availability: yes. Model weight (yolo11n.pt) not redistributed in repository; downloaded separately. Network: local inference only (no remote model call). Weight-file considerations: exclude from Git; verify checksum.')
    add_important_box(doc, 'This section provides technical compliance documentation and is not legal advice.')

    add_heading_custom(doc, '218. Current Project License Decision', 2)
    add_para(doc, 'AGPL-3.0 selected for combined repository. LICENSE file present; AGPL_COMPLIANCE.md created. No MIT claim. If architecture uses separately distributed components with different licenses, each is documented (see THIRD_PARTY_NOTICES.md). Decision is documented; legal review recommended for production release.')

    # PART XII TROUBLESHOOTING
    add_heading_custom(doc, 'Part XII — Troubleshooting', 1)
    items = [
        ('219. uvicorn: command not found', 'Use python -m uvicorn from ai-service root.'),
        ('220. ModuleNotFoundError: No module named app', 'Run from ai-service directory; use correct working directory.'),
        ('221. FastAPI health unavailable', 'Check service running; verify port 8001; check firewall.'),
        ('222. YOLO model cannot load', 'Verify yolo11n.pt present; verify checksum; verify PYTHONPATH.'),
        ('223. Model weight downloaded automatically', 'Expected if missing; verify source.'),
        ('224. Queue job remains pending', 'Start queue worker; verify database connection.'),
        ('225. Queue job fails immediately', 'Check error message; verify video asset exists; check model version.'),
        ('226. Video asset dropdown empty', 'Check valid status; verify session selected.'),
        ('227. Model-version dropdown empty', 'Check registered versions.'),
        ('228. Video asset not found', 'Check storage; verify file exists.'),
        ('229. Video file missing or unreadable', 'Re-upload; check MIME.'),
        ('230. Database row exists but file missing', 'Re-upload or delete row.'),
        ('231. Cross-service filesystem mismatch', 'Verify paths; use shared or absolute paths.'),
        ('232. Remote job ID empty', 'Wait for service response; retry.'),
        ('233. AI authentication failure', 'Check AI_SERVICE_TOKEN env setting.'),
        ('234. AI timeout', 'Increase timeout in config; verify network.'),
        ('235. Queue timeout', 'Increase timeout; verify worker.')
        ]
    for title, sol in items:
        add_heading_custom(doc, title, 3)
        add_para(doc, sol)

    # ADDITIONAL TROUBLESHOOTING (from source)
    add_heading_custom(doc, '236. Maximum Execution Time Exceeded', 3)
    add_para(doc, 'Increase max_execution_time in PHP config; verify video length; use interval > 1.')

    add_heading_custom(doc, '237. PendingRequest::method does not exist', 3)
    add_para(doc, 'Verified root cause: retry callback received PendingRequest instance (not Request object). Fixed in AiServiceClient.php by removing $request->method() call.')

    add_heading_custom(doc, '238. Unknown column analysis_jobs.deleted_at', 3)
    add_para(doc, 'Fixed by migration (softDeletes added to analysis_jobs).')

    add_heading_custom(doc, '239. Unknown column video_assets.deleted_at', 3)
    add_para(doc, 'Fixed by migration (add_deleted_at_to_video_assets_table).')

    add_heading_custom(doc, '240. Role is not displayed', 3)
    add_para(doc, 'Custom role_user pivot; verify role assignment; refresh.')

    add_heading_custom(doc, '241. 403 Forbidden', 3)
    add_para(doc, 'Policy not registered (AppServiceProvider) or role missing.')

    add_heading_custom(doc, '242. Logout option missing', 3)
    add_para(doc, 'Verify auth routes; check Blade view.')

    add_heading_custom(doc, '243. Edit/delete actions hidden', 3)
    add_para(doc, 'VideoAssetPolicy registered; @can directives present; verify user role.')

    add_heading_custom(doc, '244. Sidebar mobile issue', 3)
    add_para(doc, 'Responsive design; verify viewport.')

    add_heading_custom(doc, '245. Live preview unavailable', 3)
    add_para(doc, 'Check source; verify stream URL; verify MJPEG endpoint.')

    add_heading_custom(doc, '246. Camera disconnected or degraded', 3)
    add_para(doc, 'Check camera settings; reconnect; review health.')

    add_heading_custom(doc, '247. EZVIZ RTSP unverified', 3)
    add_warning_box(doc, 'Not verified; test with actual EZVIZ device before relying.')

    add_heading_custom(doc, '248. No events detected', 3)
    add_para(doc, 'May be valid completed result if no suspicious behavior occurred.')

    add_heading_custom(doc, '249. False or unexpected B4 event', 3)
    add_warning_box(doc, 'B4 is partial/experimental; verify temporal evidence; do not auto-penalize.')

    add_heading_custom(doc, '250. Evidence missing', 3)
    add_para(doc, 'Check evidence directory; verify authorization.')

    add_heading_custom(doc, '251. Metrics unavailable', 3)
    add_para(doc, 'Check benchmark results file.')

    add_heading_custom(doc, '252. Stale Laravel cache', 3)
    add_code_block(doc, 'php artisan optimize:clear')

    add_heading_custom(doc, '253. Worker running old code', 3)
    add_para(doc, 'Restart with queue:restart.')

    add_heading_custom(doc, '254. Migration ran but error persists', 3)
    add_para(doc, 'Clear cache; verify database; restart server.')

    add_heading_custom(doc, '255. Port already in use', 3)
    add_para(doc, 'Use different port; check active processes.')

    add_heading_custom(doc, '256. MySQL connection failure', 3)
    add_para(doc, 'Verify DB_HOST, DB_PORT, DB_DATABASE, DB_USERNAME, DB_PASSWORD.')

    add_heading_custom(doc, '257. Front-end assets missing', 3)
    add_code_block(doc, 'npm run build')

    add_heading_custom(doc, '258. Storage permissions', 3)
    add_para(doc, 'Verify directory writable by server user.')

    add_heading_custom(doc, '259. Windows path issues', 3)
    add_para(doc, 'Use correct separators; verify paths in config.')

    add_heading_custom(doc, '260. Git Bash vs PowerShell', 3)
    add_para(doc, 'Commands labeled accordingly; PowerShell uses different syntax for pipes.')

    # PART XIII PROJECT STATUS / ROADMAP
    add_heading_custom(doc, 'Part XIII — Project Status and Roadmap', 1)
    add_heading_custom(doc, '261. Final Implementation-Status Matrix', 2)
    add_table_custom(doc, ['Feature', 'Status', 'Evidence / Notes'],
        [['Recorded video mode', 'Implemented / tested', 'All modules verified'],
         ['Live surveillance mode', 'Partially implemented / partially verified', 'Webcam/test verified; RTSP unverified'],
         ['EZVIZ integration', 'Not applicable / unverified', 'Not verified'],
         ['Person detection', 'Implemented / tested', 'YOLO11n benchmark'],
         ['Phone detection', 'Implemented / tested', 'Class 66 annotated'],
         ['Tracking', 'Implemented', 'SimpleCentroidTracker'],
         ['Orientation', 'Experimental', 'Geometric-v1; not validated'],
         ['Temporal events B1-B3', 'Partially implemented', 'Rules exist; synthetic only'],
         ['Temporal event B4', 'Experimental', 'Partial; not fully validated'],
         ['Dashboard', 'Implemented / tested', 'Full CRUD verified'],
         ['Evidence', 'Implemented', 'Protected directory'],
         ['Human review', 'Implemented', 'Required; not automated'],
         ['Security tests', 'Partial', 'Audit done; no penetration test'],
         ['Dataset', 'Blocked (real) / active (synthetic)', 'MANIFEST.json v0.1.0'],
         ['Training', 'Partial', 'Benchmark only'],
         ['Evaluation', 'Partial', 'Synthetic only; real blocked'],
         ['License', 'AGPL-3.0 selected', 'Compliance docs started'],
         ['CI / Dependabot', 'Blocked / missing', 'Not configured'],
         ['Production readiness', 'Blocked', 'Release checklist incomplete']])

    add_heading_custom(doc, '262–268. Status Groupings', 2)
    add_para(doc, 'Implemented and runtime validated: recorded mode, detection, dashboard, audit, evidence, tests.')
    add_para(doc, 'Implemented but partially tested: live mode (RTSP unverified); temporal rules (synthetic only).')
    add_para(doc, 'Partially implemented: orientation, B4 event, training.')
    add_para(doc, 'Experimental: geometric orientation, temporal model (if data justifies).')
    add_para(doc, 'Blocked: real participant evaluation; CI/dependabot; production deployment.')
    add_para(doc, 'Not applicable: EZVIZ direct integration (unverified).')

    add_heading_custom(doc, '269. Recommended Future Work', 2)
    recommendations = [
        'Approved staged dataset with verified consent.',
        'Real-world event evaluation (after consent verified).',
        'Improved multi-person orientation.',
        'Robust seat-region configuration.',
        'Verified EZVIZ integration.',
        'GPU comparison.',
        'CI / Dependabot.',
        'Vulnerability scanning.',
        'Production deployment (only after all criteria met).',
    ]
    for r in recommendations:
        p = doc.add_paragraph('• ' + r, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.name = 'Calibri'; run.font.size = Pt(10)

    add_heading_custom(doc, '270. Release-Readiness Statement', 2)
    add_warning_box(doc, 'Not production-ready. Release acceptance criteria not met (CI/dependabot missing; user acceptance not fully signed off; full static analysis not executed; GPU benchmark unavailable). No version tag pushed. Remediation list documented in REMEDIATION_REPORT.md.')

    # PART XIV APPENDICES
    add_heading_custom(doc, 'Part XIV — Appendices', 1)
    add_heading_custom(doc, '271. Complete Command Reference', 2)
    add_code_block(doc, '# Repository\ngit status\ngit log --oneline -5\n\n# AI service\npython -m uvicorn app.main:app --host 127.0.0.1 --port 8001\n\n# Laravel\nphp artisan serve\nphp artisan queue:work --queue=default --tries=1 --timeout=900 -v\nphp artisan test\nphp artisan migrate --force\nphp artisan optimize:clear\n\n# Tests\nphp artisan test tests/Feature/VideoAssetFailureTest.php\n')

    add_heading_custom(doc, '272. Environment-Variable Reference', 2)
    add_table_custom(doc, ['Key', 'Purpose', 'Required'],
        [['APP_KEY', 'Encryption', 'Yes'],
         ['DB_CONNECTION', 'DB type', 'Yes'],
         ['DB_HOST', 'DB host', 'Yes'],
         ['DB_DATABASE', 'DB name', 'Yes'],
         ['DB_USERNAME', 'DB user', 'Yes'],
         ['DB_PASSWORD', 'DB password', 'Yes'],
         ['AI_SERVICE_TOKEN', 'API auth', 'Yes'],
         ['AI_SERVICE_BASE_URL', 'AI URL', 'No'],
         ['APP_ENV', 'Environment', 'No'],
         ['APP_DEBUG', 'Debug', 'No (false in prod)'],
         ['SESSION_ENCRYPT', 'Session', 'No (true in prod)']])

    add_heading_custom(doc, '273. Role and Permission Matrix', 2)
    add_table_custom(doc, ['Role', 'Edit Video', 'Del Video', 'Manage Job', 'Review', 'Audit'],
        [['system_admin', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
         ['exam_admin', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
         ['invigilator', 'No', 'No', 'View', 'No', 'No'],
         ['reviewer', 'No', 'No', 'View', 'Yes', 'No'],
         ['auditor', 'No', 'No', 'View', 'No', 'Yes']])

    add_heading_custom(doc, '274. Event Taxonomy Matrix', 2)
    add_table_custom(doc, ['Category', 'Code', 'Meaning', 'Status'],
        [['Detection', 'D1', 'Person', 'Implemented'],
         ['Detection', 'D2', 'Phone', 'Implemented'],
         ['Behavior', 'B1-B3', 'Looking', 'Partial'],
         ['Behavior', 'B4', 'Seat departure', 'Experimental'],
         ['State', 'S1-S2', 'Normal/Insufficient', 'Implemented']])

    add_heading_custom(doc, '275. Job-State Transition Matrix', 2)
    add_table_custom(doc, ['From', 'To', 'Trigger'],
        [['Pending', 'Queued', 'Worker picks'],
         ['Queued', 'Processing', 'Worker starts'],
         ['Processing', 'Completed', 'Success'],
         ['Processing', 'Failed', 'Error'],
         ['Failed', 'Queued', 'Retry'],
         ['Queued/Processing', 'Cancelled', 'User cancel'],
         ['Cancelled', 'Queued', 'Retry']])

    add_heading_custom(doc, '276. API Endpoint Summary', 2)
    add_para(doc, 'See Part V, Section 75.')

    add_heading_custom(doc, '277. Database Table Summary', 2)
    add_para(doc, 'See Part V, Section 70.')

    add_heading_custom(doc, '278. File and Directory Reference', 2)
    add_para(doc, 'Key directories: ai-service/, dashboard/, docs/, research/, storage/, scripts/, .gitignored datasets/ and evidence/.')

    add_heading_custom(doc, '279. Audit-Event Reference', 2)
    add_para(doc, 'Video asset actions (upload/update/delete/blocked), job actions (create/approve/review/delete/retry), model actions.')

    add_heading_custom(doc, '280. Error / Status Code Reference', 2)
    add_table_custom(doc, ['Code', 'Meaning'],
        [['200', 'Success'], ['201', 'Created'], ['400', 'Validation'], ['401', 'Auth'], ['403', 'Forbidden'], ['404', 'Not found'], ['500', 'Server error'], ['503', 'Service unavailable']])

    add_heading_custom(doc, '281. Responsible-Use Checklist', 2)
    checklist_items = [
        'AI alerts require human review.',
        'No automated disciplinary decision.',
        'Evidence protected.',
        'Audit log complete.',
        'Consent verified for real recordings.',
        'Synthetic/non-identifiable only for current evaluation.',
    ]
    for item in checklist_items:
        p = doc.add_paragraph('☐  ' + item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.name = 'Calibri'; run.font.size = Pt(9)

    add_heading_custom(doc, '282. Operator Daily Checklist', 2)
    daily = [
        'Start services (FastAPI, Laravel, worker).',
        'Check dashboard health.',
        'Review pending jobs.',
        'Review evidence / alerts.',
        'Check audit logs.',
        'Verify storage space.',
    ]
    for item in daily:
        p = doc.add_paragraph('• ' + item)
        p.paragraph_format.space_after = Pt(2)

    add_heading_custom(doc, '283. Administrator Maintenance Checklist', 2)
    maint = [
        'Clear caches when needed.',
        'Restart worker after updates.',
        'Backup database.',
        'Review retention actions.',
        'Check failed jobs.',
        'Update model versions with verified checksums.',
    ]
    for item in maint:
        p = doc.add_paragraph('• ' + item)
        p.paragraph_format.space_after = Pt(2)

    add_heading_custom(doc, '284. Developer Release Checklist', 2)
    dev = [
        'Verify tests pass.',
        'No secrets committed.',
        'Documentation updated.',
        'Model checksum verified.',
        'Dataset manifest updated if needed.',
        'No raw videos or weights in commit.',
    ]
    for item in dev:
        p = doc.add_paragraph('• ' + item)
        p.paragraph_format.space_after = Pt(2)

    add_heading_custom(doc, '285. References', 2)
    add_para(doc, 'See DATA_PLAN.md, THREAT_MODEL.md, MODEL_BASELINE.md, BENCHMARKING.md, and repository documentation for full references.')

    # Save
    output_path = 'AI_Classroom_Cheating_Detection_System_Complete_Documentation.docx'
    doc.save(output_path)
    print(f"DOCX saved: {output_path}")

if __name__ == '__main__':
    main()
